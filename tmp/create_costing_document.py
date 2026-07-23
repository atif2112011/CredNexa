from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from pathlib import Path

OUT = Path(r"D:\Practise Projects\EMI Shield\output\EMI_Shield_Platform_Cost_Estimate.docx")

INK = "17324D"
BLUE = "2E74B5"
MUTED = "5C6B7A"
FILL = "E8EEF5"
LIGHT = "F4F6F9"
GREEN = "1F6B45"

def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)

def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")

def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = OxmlElement("w:tblInd")
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_ind)
    grid = table._tbl.tblGrid
    for col, width in zip(grid.gridCol_lst, widths):
        col.set(qn("w:w"), str(width))
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            set_cell_width(cell, width)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            tc_pr = cell._tc.get_or_add_tcPr()
            mar = tc_pr.first_child_found_in("w:tcMar")
            if mar is None:
                mar = OxmlElement("w:tcMar")
                tc_pr.append(mar)
            for side in ("top", "start", "bottom", "end"):
                node = mar.find(qn(f"w:{side}"))
                if node is None:
                    node = OxmlElement(f"w:{side}")
                    mar.append(node)
                node.set(qn("w:w"), "100" if side in ("top", "bottom") else "120")
                node.set(qn("w:type"), "dxa")

def set_font(run, size=11, color=INK, bold=False, italic=False):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic

def add_text(p, text, **kwargs):
    r = p.add_run(text)
    set_font(r, **kwargs)
    return r

def set_para(p, before=0, after=6, line=1.1):
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line

def cell_text(cell, text, bold=False, color=INK, size=10):
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_para(p, after=0, line=1.08)
    add_text(p, str(text), size=size, color=color, bold=bold)

def add_table(doc, headers, rows, widths, money_cols=()):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    for i, h in enumerate(headers):
        cell_text(table.rows[0].cells[i], h, bold=True, color=INK, size=10)
        set_cell_shading(table.rows[0].cells[i], FILL)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cell_text(cells[i], value, bold=False, size=10)
            if i in money_cols:
                cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table

def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    set_para(p, before=14 if level == 1 else 9, after=6)
    add_text(p, text, size=16 if level == 1 else 13, color=BLUE if level < 3 else INK, bold=True)
    return p

def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    set_para(p, after=4, line=1.12)
    for run in p.runs:
        set_font(run, size=11)
    add_text(p, text, size=11)

doc = Document()
sec = doc.sections[0]
sec.top_margin = Inches(0.85)
sec.bottom_margin = Inches(0.8)
sec.left_margin = Inches(0.85)
sec.right_margin = Inches(0.85)
sec.header_distance = Inches(0.3)
sec.footer_distance = Inches(0.35)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
normal.font.size = Pt(11)

header = sec.header.paragraphs[0]
header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
set_para(header, after=0)
add_text(header, "EMI Shield | Internal costing estimate", size=8.5, color=MUTED)
footer = sec.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_para(footer, after=0)
add_text(footer, "Planning estimate - update with actual Cloud Billing and OTP provider invoices", size=8.5, color=MUTED)

p = doc.add_paragraph()
set_para(p, after=3)
add_text(p, "EMI SHIELD", size=10, color=BLUE, bold=True)
p = doc.add_paragraph()
set_para(p, after=4)
add_text(p, "Platform Cost Estimate", size=25, color=INK, bold=True)
p = doc.add_paragraph()
set_para(p, after=14)
add_text(p, "Scale assumption: 5 partners | 50 total tenants | 10,000 financed devices", size=12, color=MUTED)

add_heading(doc, "Executive summary")
p = doc.add_paragraph()
set_para(p, after=8)
add_text(p, "Known onboarding OTP cost: ", size=11, bold=True)
add_text(p, "Rs. 2,011", size=11, color=GREEN, bold=True)
add_text(p, " for 10,055 OTPs at Rs. 0.20 per message.", size=11)

summary_rows = [
    ("One-time OTP onboarding", "Rs. 2,011", "Known; 10,055 messages"),
    ("Recurring OTP - recovery only", "Rs. 0.20 / month", "One partner/tenant password reset assumed per month"),
    ("MongoDB Atlas Free (M0)", "Rs. 0 / month", "Selected for this estimate; 512 MB shared free tier"),
    ("Firebase/Google services budget", "Rs. 1,500 / month", "Planning reserve; monitor actual billing"),
    ("Estimated monthly platform cost", "Rs. 1,500.20", "Atlas Free + Firebase budget + recovery OTP; excludes other vendors"),
    ("First-year core budget", "Rs. 20,013", "One-time onboarding + 12 months of estimated monthly platform cost")
]
add_table(doc, ["Cost area", "Estimate", "Basis"], summary_rows, [2850, 1800, 4710], money_cols=(1,))

add_heading(doc, "1. Volume assumptions")
volume_rows = [
    ("Partners", "5"),
    ("Total tenants", "50"),
    ("Average tenants per partner", "10"),
    ("Devices per tenant", "200"),
    ("Total borrower devices", "10,000")
]
add_table(doc, ["Metric", "Volume"], volume_rows, [5700, 3660])

add_heading(doc, "2. One-time onboarding OTP cost")
p = doc.add_paragraph()
set_para(p, after=6)
add_text(p, "Assumption: one OTP is sent for each partner, tenant administrator, and borrower-device onboarding.", size=10.5, color=MUTED, italic=True)
otp_rows = [
    ("Partner onboarding", "5", "1", "5", "Rs. 1"),
    ("Tenant onboarding", "50", "1", "50", "Rs. 10"),
    ("Borrower device onboarding", "10,000", "1", "10,000", "Rs. 2,000"),
    ("Total", "", "", "10,055", "Rs. 2,011")
]
table = add_table(doc, ["OTP use", "Users/devices", "OTPs each", "Messages", "Cost"], otp_rows, [3100, 1400, 1400, 1600, 1860], money_cols=(4,))
for cell in table.rows[-1].cells:
    set_cell_shading(cell, LIGHT)
    for run in cell.paragraphs[0].runs:
        run.bold = True

add_heading(doc, "3. Recurring OTP cost - password recovery only")
p = doc.add_paragraph()
set_para(p, after=6)
add_text(p, "Partners, tenants, and borrowers log in only once. The platform should keep the user session active after the first login. The following recovery-only estimate applies to partner and tenant password resets; borrower re-enrolment/reinstallation is excluded.", size=10.5, color=MUTED, italic=True)
scenario_rows = [
    ("Low", "0", "0", "0", "Rs. 0"),
    ("Base", "1", "1", "1", "Rs. 0.20"),
    ("High", "6", "6", "6", "Rs. 1.20")
]
add_table(doc, ["Scenario", "Password resets/month", "Recovery OTPs", "Total OTPs", "Monthly cost"], scenario_rows, [1300, 2550, 1700, 1800, 2010], money_cols=(4,))

add_heading(doc, "4. Firebase Functions and related services")
p = doc.add_paragraph()
set_para(p, after=6)
add_text(p, "The backend is deployed as Firebase Functions v2 in asia-south1 and includes six scheduled functions. Its final bill depends on execution time, CPU/memory allocation, MongoDB response time, outbound traffic, logs, and storage - not only device count.", size=10.5, color=MUTED, italic=True)

firebase_rows = [
    ("API traffic planning case", "4 calls/device/day x 10,000 devices x 30 days = 1.2 million API requests/month"),
    ("Scheduled functions", "Approx. 50,460 executions/month across six jobs; Scheduler charge is about Rs. 29/month after 3 free jobs"),
    ("Function requests", "Likely within the 2 million monthly free invocation tier in this planning case; compute time can still be billable"),
    ("FCM notifications", "Rs. 0 (Firebase Cloud Messaging is no-cost)"),
    ("Cloud Storage", "Rs. 0 while within the included 5 GB, then usage-based"),
    ("Planning reserve", "Rs. 1,500/month for functions, scheduler, storage/egress, and logging variability")
]
add_table(doc, ["Service", "Planning treatment"], firebase_rows, [2650, 6710])

add_heading(doc, "5. Monthly and first-year planning budget")
budget_rows = [
    ("OTP service - recovery-only base case", "Rs. 0.20", "Rs. 2.40"),
    ("MongoDB Atlas Free (M0)", "Rs. 0", "Rs. 0"),
    ("Firebase/Google services reserve", "Rs. 1,500", "Rs. 18,000"),
    ("Estimated monthly platform cost", "Rs. 1,500.20", "Rs. 18,002.40"),
    ("One-time onboarding OTP", "Rs. 2,011", "Rs. 2,011"),
    ("First-year core budget", "", "Rs. 20,013.40")
]
table = add_table(doc, ["Cost area", "Monthly", "Year 1"], budget_rows, [4200, 2500, 2660], money_cols=(1,2))
for cell in table.rows[-1].cells:
    set_cell_shading(cell, LIGHT)
    for run in cell.paragraphs[0].runs:
        run.bold = True

add_heading(doc, "6. Not included in this estimate")
for text in [
    "MongoDB Atlas backup or upgrade charges, database data-transfer charges, and any Atlas paid-tier costs if the free 512 MB limit is exceeded.",
    "Custom domain, email provider, payment gateway/QR settlement fees, Play Console fees, and KYC/Aadhaar provider fees.",
    "Engineering, support staff, tenant training, customer support, legal/compliance, GST, and contingency outside the Firebase reserve.",
    "Costs caused by higher device sync frequency, large payment-proof uploads, app APK downloads, or unusually long scheduled jobs."
]:
    add_bullet(doc, text)

add_heading(doc, "7. Cost-control recommendations")
for text in [
    "Keep the first-login session and refresh token active so partner, tenant, and borrower users do not receive another OTP during normal app use.",
    "Use OTP only for login, recovery, onboarding, and high-risk actions; apply rate limits and cooldowns.",
    "Track Cloud Functions v2 CPU seconds, memory seconds, egress, Cloud Scheduler, Storage, Logging, and MongoDB separately from day one.",
    "Set Cloud Billing budgets at Rs. 1,500 (alert), Rs. 2,500 (review), and Rs. 5,000 (urgent review) per month for Firebase/Google services."
]:
    add_bullet(doc, text)

add_heading(doc, "Pricing references and assumptions")
p = doc.add_paragraph()
set_para(p, after=4)
add_text(p, "OTP rate supplied by business: Rs. 0.20 per SMS. Currency conversion used only for planning reserve: Rs. 96 per USD, rounded from July 2026 RBI reference-rate information.", size=9.5, color=MUTED)
for source in [
    "Firebase pricing: https://firebase.google.com/pricing",
    "Firebase Blaze plan: https://firebase.google.com/docs/projects/billing/firebase-pricing-plans",
    "Cloud Run functions pricing: https://cloud.google.com/functions/pricing-1stgen",
    "Cloud Scheduler pricing: https://cloud.google.com/scheduler/pricing",
    "MongoDB Atlas Free tier limits: https://www.mongodb.com/docs/atlas/reference/free-shared-limitations/"
]:
    p = doc.add_paragraph()
    set_para(p, after=2)
    add_text(p, source, size=9.5, color=MUTED)

OUT.parent.mkdir(parents=True, exist_ok=True)
doc.core_properties.title = "EMI Shield Platform Cost Estimate"
doc.core_properties.subject = "Cost estimate for 5 partners, 50 tenants, and 10,000 financed devices"
doc.core_properties.author = "EMI Shield"
doc.save(OUT)
print(OUT)
